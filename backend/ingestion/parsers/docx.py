"""DOCX parser."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from ingestion.models import DocumentFormat, ParsedDocument, SourceDocument
from ingestion.parsers.exceptions import ParserError
from ingestion.parsers.utils import build_parser_metadata, normalize_newlines


class DocxParser:
    """Extract normalized text from DOCX paragraphs and tables."""

    name = "docx"
    version = "1.0"
    supported_formats = (DocumentFormat.DOCX,)

    def parse(self, source: SourceDocument, content: bytes) -> ParsedDocument:
        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise ParserError(f"Unable to read DOCX content: {exc}") from exc

        paragraph_text = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        table_text: list[str] = []
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    table_text.append(" | ".join(cells))

        text_parts = paragraph_text + table_text
        if not text_parts:
            raise ParserError("DOCX does not contain extractable text.")

        return ParsedDocument(
            id=f"parsed-{source.id}",
            source_document_id=source.id,
            text_content=normalize_newlines("\n\n".join(text_parts)),
            parser_name=self.name,
            parser_version=self.version,
            parser_metadata=build_parser_metadata(
                paragraph_count=len(paragraph_text),
                table_row_count=len(table_text),
            ),
        )