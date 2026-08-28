"""Unit tests for concrete local document parsers."""

from __future__ import annotations

import csv
import importlib.util
from io import BytesIO

import pytest
from docx import Document
from openpyxl import Workbook

from ingestion.models import (
    DocumentFormat,
    ParsedDocument,
    ParserWarning,
    SourceDocument,
    SourceType,
)
from ingestion.parsers.csv import CsvParser
from ingestion.parsers.docx import DocxParser
from ingestion.parsers.exceptions import ParserError
from ingestion.parsers.html import HtmlParser
from ingestion.parsers.json import JsonParser
from ingestion.parsers.pdf import PdfParser
from ingestion.parsers.txt import TextParser
from ingestion.parsers.xlsx import XlsxParser


def _source(document_id: str, document_format: DocumentFormat) -> SourceDocument:
    return SourceDocument(
        id=document_id,
        source_type=SourceType.FILE_UPLOAD,
        document_format=document_format,
        filename=f"sample.{document_format.value}",
    )


def test_text_parser_decodes_and_normalizes() -> None:
    parser = TextParser()
    parsed = parser.parse(_source("doc-txt", DocumentFormat.TXT), b"hello\r\nworld")

    assert parsed.text_content == "hello\nworld"
    assert parsed.parser_metadata["encoding"] == "utf-8"


def test_json_parser_creates_records_for_array_of_objects() -> None:
    parser = JsonParser()
    parsed = parser.parse(
        _source("doc-json", DocumentFormat.JSON),
        b'[{"claim_id": "1"}, {"claim_id": "2"}]',
    )

    assert len(parsed.records) == 2
    assert parsed.records[1].fields["claim_id"] == "2"


def test_json_parser_rejects_invalid_json() -> None:
    parser = JsonParser()
    with pytest.raises(ParserError, match="Invalid JSON"):
        parser.parse(_source("doc-json", DocumentFormat.JSON), b"{bad json")


def test_csv_parser_creates_structured_records() -> None:
    parser = CsvParser()
    parsed = parser.parse(
        _source("doc-csv", DocumentFormat.CSV),
        b"claim_id,amount\n1,100\n2,250\n",
    )

    assert len(parsed.records) == 2
    assert parsed.records[0].fields["claim_id"] == "1"
    assert parsed.parser_metadata["has_header"] is True


def test_csv_parser_rejects_empty_content() -> None:
    parser = CsvParser()
    with pytest.raises(ParserError, match="empty"):
        parser.parse(_source("doc-csv", DocumentFormat.CSV), b"   ")


def test_csv_parser_falls_back_when_dialect_sniffing_fails(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    def fail_sniff(
        self: object,
        sample: str,
        delimiters: str | None = None,
    ) -> csv.Dialect:
        raise csv.Error("cannot sniff")

    monkeypatch.setattr("csv.Sniffer.sniff", fail_sniff)

    parser = CsvParser()
    parsed = parser.parse(
        _source("doc-csv", DocumentFormat.CSV),
        b"claim_id,amount\n1,100\n",
    )

    assert parsed.records[0].fields["claim_id"] == "1"
    assert parsed.parser_metadata["delimiter"] == ","


def test_csv_parser_generates_column_names_without_headers(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    def no_header(self: object, sample: str) -> bool:
        del self, sample
        return False

    monkeypatch.setattr("csv.Sniffer.has_header", no_header)

    parser = CsvParser()
    parsed = parser.parse(
        _source("doc-csv", DocumentFormat.CSV),
        b"1,100\n2,250\n",
    )

    assert parsed.records[0].fields == {"column_1": "1", "column_2": "100"}
    assert parsed.parser_metadata["has_header"] is False


def test_csv_parser_rejects_no_rows_after_reader(monkeypatch: pytest.MonkeyPatch) -> None:
    def empty_reader(
        csvfile: object,
        dialect: csv.Dialect | str = "excel",
        **fmtparams: object,
    ) -> list[list[str]]:
        del csvfile, dialect, fmtparams
        return []

    monkeypatch.setattr("csv.reader", empty_reader)

    parser = CsvParser()
    with pytest.raises(ParserError, match="does not contain any rows"):
        parser.parse(_source("doc-csv", DocumentFormat.CSV), b"claim_id\n1\n")


def test_csv_parser_rejects_headers_without_data_rows() -> None:
    parser = CsvParser()
    with pytest.raises(ParserError, match="headers but no data rows"):
        parser.parse(_source("doc-csv", DocumentFormat.CSV), b"claim_id,amount\n")


def test_html_parser_extracts_visible_text() -> None:
    parser = HtmlParser()
    parsed = parser.parse(
        _source("doc-html", DocumentFormat.HTML),
        b"""
        <html>
          <head><title>Ignored</title><script>alert("x")</script></head>
          <body><h1>Claim Summary</h1><p>Claim ID C-1 &amp; amount $42.</p></body>
        </html>
        """,
    )

    assert parsed.text_content == "# Claim Summary\n\nClaim ID C-1 & amount $42."
    assert parsed.parser_metadata["encoding"] == "utf-8"


def test_xlsx_parser_creates_records_with_sheet_metadata() -> None:
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "Claims"
    sheet.append(["claim_id", "amount"])
    sheet.append(["C-1", 42])
    output = BytesIO()
    workbook.save(output)

    parser = XlsxParser()
    parsed = parser.parse(_source("doc-xlsx", DocumentFormat.XLSX), output.getvalue())

    assert len(parsed.records) == 1
    assert parsed.records[0].fields["claim_id"] == "C-1"
    assert parsed.records[0].metadata["sheet_name"] == "Claims"


def test_xlsx_parser_rejects_invalid_workbook() -> None:
    parser = XlsxParser()
    with pytest.raises(ParserError, match="Unable to read XLSX"):
        parser.parse(_source("doc-xlsx", DocumentFormat.XLSX), b"not an xlsx file")


def test_docx_parser_extracts_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("Paragraph text")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    output = BytesIO()
    document.save(output)

    parser = DocxParser()
    parsed = parser.parse(_source("doc-docx", DocumentFormat.DOCX), output.getvalue())

    assert "Paragraph text" in (parsed.text_content or "")
    assert "A | B" in (parsed.text_content or "")
    assert parsed.parser_metadata["table_row_count"] == 1


def test_docx_parser_rejects_invalid_docx() -> None:
    parser = DocxParser()
    with pytest.raises(ParserError, match="Unable to read DOCX"):
        parser.parse(_source("doc-docx", DocumentFormat.DOCX), b"bad docx")


class _FakePage:
    def __init__(self, text: str) -> None:
        self._text = text

    def extract_text(self) -> str:
        return self._text


class _FakePdfReader:
    def __init__(self, _content: BytesIO) -> None:
        self.is_encrypted = False
        self.pages = [_FakePage("First page"), _FakePage("Second page")]


def test_pdf_parser_extracts_text(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("ingestion.parsers.pdf.PdfReader", _FakePdfReader)

    parser = PdfParser()
    parsed = parser.parse(_source("doc-pdf", DocumentFormat.PDF), b"fake pdf")

    assert parsed.text_content == "First page\n\nSecond page"
    assert parsed.parser_metadata["page_count"] == 2


class _EncryptedPdfReader:
    def __init__(self, _content: BytesIO) -> None:
        self.is_encrypted = True
        self.pages = []


def test_pdf_parser_rejects_encrypted_files(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("ingestion.parsers.pdf.PdfReader", _EncryptedPdfReader)

    parser = PdfParser()
    with pytest.raises(ParserError, match="Encrypted PDF"):
        parser.parse(_source("doc-pdf", DocumentFormat.PDF), b"fake pdf")


# --- HTML structural fidelity (ingestion.02) ----------------------------------


def test_html_parser_preserves_heading_markers() -> None:
    parser = HtmlParser()
    parsed = parser.parse(
        _source("doc-html", DocumentFormat.HTML),
        b"<body><h1>Title</h1><h2>Section</h2><h3>Sub</h3><p>Body.</p></body>",
    )

    assert parsed.text_content == "# Title\n\n## Section\n\n### Sub\n\nBody."
    assert parsed.parser_metadata["heading_count"] == 3


def test_html_parser_preserves_link_targets() -> None:
    parser = HtmlParser()
    parsed = parser.parse(
        _source("doc-html", DocumentFormat.HTML),
        b'<body><p>See <a href="https://example.com/x">the policy</a> now.</p></body>',
    )

    assert parsed.text_content == "See [the policy](https://example.com/x) now."
    assert parsed.parser_metadata["link_count"] == 1


def test_html_parser_keeps_anchor_text_when_href_missing() -> None:
    parser = HtmlParser()
    parsed = parser.parse(
        _source("doc-html", DocumentFormat.HTML),
        b"<body><p>Plain <a>anchor</a> text.</p></body>",
    )

    assert parsed.text_content == "Plain anchor text."
    assert parsed.parser_metadata["link_count"] == 0


def test_html_parser_renders_table_as_markdown() -> None:
    parser = HtmlParser()
    parsed = parser.parse(
        _source("doc-html", DocumentFormat.HTML),
        b"""
        <body>
          <table>
            <tr><th>Provider</th><th>NPI</th></tr>
            <tr><td>Acme</td><td>123</td></tr>
            <tr><td>Beta</td><td>456</td></tr>
          </table>
        </body>
        """,
    )

    assert parsed.text_content == (
        "| Provider | NPI |\n"
        "| --- | --- |\n"
        "| Acme | 123 |\n"
        "| Beta | 456 |"
    )
    assert parsed.parser_metadata["table_count"] == 1


def test_html_parser_pads_ragged_table_rows() -> None:
    parser = HtmlParser()
    parsed = parser.parse(
        _source("doc-html", DocumentFormat.HTML),
        b"<body><table><tr><th>A</th><th>B</th></tr><tr><td>only</td></tr></table></body>",
    )

    assert parsed.text_content == "| A | B |\n| --- | --- |\n| only |  |"


def test_html_parser_flattens_nested_tables_into_parent_cell() -> None:
    parser = HtmlParser()
    parsed = parser.parse(
        _source("doc-html", DocumentFormat.HTML),
        b"""
        <body>
          <table>
            <tr><th>Outer</th></tr>
            <tr><td>cell <table><tr><td>inner1</td><td>inner2</td></tr></table></td></tr>
          </table>
        </body>
        """,
    )

    # Nested table flattened into the parent cell; outer table stays valid markdown.
    assert parsed.text_content == "| Outer |\n| --- |\n| cell inner1 inner2 |"
    assert parsed.parser_metadata["table_count"] == 2


def test_html_parser_counts_all_structures_in_metadata() -> None:
    parser = HtmlParser()
    parsed = parser.parse(
        _source("doc-html", DocumentFormat.HTML),
        b"""
        <body>
          <h1>Report</h1>
          <p>Visit <a href="https://a.test">A</a> and <a href="https://b.test">B</a>.</p>
          <table><tr><th>H</th></tr><tr><td>v</td></tr></table>
        </body>
        """,
    )

    assert parsed.parser_metadata["heading_count"] == 1
    assert parsed.parser_metadata["link_count"] == 2
    assert parsed.parser_metadata["table_count"] == 1
    assert (parsed.text_content or "").startswith("# Report")
    assert "[A](https://a.test)" in (parsed.text_content or "")


# --- Typed parser warnings (ingestion.24) -------------------------------------


def _codes(parsed: ParsedDocument) -> set[str]:
    return {warning.code for warning in parsed.warnings}


def test_text_parser_warns_on_charset_fallback() -> None:
    parser = TextParser()
    parsed = parser.parse(_source("doc-txt", DocumentFormat.TXT), b"caf\xe9 latte")

    assert parsed.text_content == "café latte"
    assert "text.charset_fallback" in _codes(parsed)
    assert parsed.warnings[0].severity == "info"


def test_text_parser_emits_no_warnings_for_clean_utf8() -> None:
    parser = TextParser()
    parsed = parser.parse(_source("doc-txt", DocumentFormat.TXT), b"plain ascii")

    assert parsed.warnings == []


def test_html_parser_warns_on_charset_fallback() -> None:
    parser = HtmlParser()
    parsed = parser.parse(
        _source("doc-html", DocumentFormat.HTML),
        b"<html><body><p>caf\xe9</p></body></html>",
    )

    assert "html.charset_fallback" in _codes(parsed)


def test_json_parser_warns_on_heterogeneous_array() -> None:
    parser = JsonParser()
    parsed = parser.parse(
        _source("doc-json", DocumentFormat.JSON),
        b'[1, {"claim_id": "1"}]',
    )

    assert parsed.records == []
    assert "json.heterogeneous_array" in _codes(parsed)


def test_json_parser_warns_on_scalar_root() -> None:
    parser = JsonParser()
    parsed = parser.parse(_source("doc-json", DocumentFormat.JSON), b"42")

    assert "json.scalar_root" in _codes(parsed)


def test_csv_parser_warns_on_ragged_row() -> None:
    parser = CsvParser()
    parsed = parser.parse(
        _source("doc-csv", DocumentFormat.CSV),
        b"claim_id,amount\n1,100,extra\n",
    )

    ragged = [w for w in parsed.warnings if w.code == "csv.ragged_row"]
    assert len(ragged) == 1
    assert ragged[0].row_index == 0
    # Extra cell dropped; declared columns preserved.
    assert parsed.records[0].fields == {"claim_id": "1", "amount": "100"}


def test_csv_parser_warns_on_dialect_fallback(monkeypatch: pytest.MonkeyPatch) -> None:
    def fail_sniff(self: object, sample: str, delimiters: str | None = None) -> csv.Dialect:
        del self, sample, delimiters
        raise csv.Error("cannot sniff")

    monkeypatch.setattr("csv.Sniffer.sniff", fail_sniff)

    parser = CsvParser()
    parsed = parser.parse(_source("doc-csv", DocumentFormat.CSV), b"claim_id,amount\n1,100\n")

    assert "csv.dialect_fallback" in _codes(parsed)


def test_xlsx_parser_warns_on_blank_row_skipped() -> None:
    workbook = Workbook()
    sheet = workbook.active
    assert sheet is not None
    sheet.title = "Claims"
    sheet.append(["claim_id", "amount"])
    sheet.append(["C-1", 42])
    sheet.append([None, None])
    sheet.append(["C-2", 43])
    output = BytesIO()
    workbook.save(output)

    parser = XlsxParser()
    parsed = parser.parse(_source("doc-xlsx", DocumentFormat.XLSX), output.getvalue())

    assert len(parsed.records) == 2
    blank = [w for w in parsed.warnings if w.code == "xlsx.blank_row_skipped"]
    assert len(blank) == 1
    assert blank[0].row_index == 2


def test_pdf_parser_warns_on_empty_page(monkeypatch: pytest.MonkeyPatch) -> None:
    class _PartiallyEmptyReader:
        def __init__(self, _content: BytesIO) -> None:
            self.is_encrypted = False
            self.pages = [_FakePage("First page"), _FakePage("   ")]

    monkeypatch.setattr("ingestion.parsers.pdf.PdfReader", _PartiallyEmptyReader)

    parser = PdfParser()
    parsed = parser.parse(_source("doc-pdf", DocumentFormat.PDF), b"fake pdf")

    empty = [w for w in parsed.warnings if w.code == "pdf.empty_page"]
    assert len(empty) == 1
    assert empty[0].page_number == 2


def test_docx_parser_warns_on_empty_paragraphs() -> None:
    document = Document()
    document.add_paragraph("   ")
    document.add_paragraph("Real content")
    output = BytesIO()
    document.save(output)

    parser = DocxParser()
    parsed = parser.parse(_source("doc-docx", DocumentFormat.DOCX), output.getvalue())

    assert "docx.empty_paragraph_skipped" in _codes(parsed)


def test_parser_warning_round_trips_through_serialization() -> None:
    parser = PdfParser()  # any parser; reuse the empty-page warning path

    class _Reader:
        def __init__(self, _content: BytesIO) -> None:
            self.is_encrypted = False
            self.pages = [_FakePage("kept"), _FakePage("")]

    import ingestion.parsers.pdf as pdf_module

    original = pdf_module.PdfReader
    pdf_module.PdfReader = _Reader  # type: ignore[assignment]
    try:
        parsed = parser.parse(_source("doc-pdf", DocumentFormat.PDF), b"fake pdf")
    finally:
        pdf_module.PdfReader = original  # type: ignore[assignment]

    restored = ParsedDocument.model_validate_json(parsed.model_dump_json())
    assert len(restored.warnings) == 1
    assert restored.warnings[0].code == "pdf.empty_page"
    assert restored.warnings[0].page_number == 2
    assert restored.warnings[0].severity == "warning"


# --- PDF OCR fallback (ingestion.03) ------------------------------------------


def _pdf_reader_with_pages(page_texts: list[str]) -> type:
    class _Reader:
        def __init__(self, _content: BytesIO) -> None:
            self.is_encrypted = False
            self.pages = [_FakePage(text) for text in page_texts]

    return _Reader


class _StubOcrAdapter:
    """In-tree OCR stub: returns canned text per 1-based page; records calls."""

    def __init__(self, page_text: dict[int, str] | None = None) -> None:
        self.calls: list[int] = []
        self._page_text = page_text or {}

    def recognize_page(self, content: bytes, page_number: int) -> str:
        del content
        self.calls.append(page_number)
        return self._page_text.get(page_number, "")


def test_pdf_parser_ocrs_image_only_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("ingestion.parsers.pdf.PdfReader", _pdf_reader_with_pages(["", ""]))
    adapter = _StubOcrAdapter({1: "Recognized one", 2: "Recognized two"})

    parser = PdfParser(ocr_adapter=adapter)
    parsed = parser.parse(_source("doc-pdf", DocumentFormat.PDF), b"image pdf")

    assert parsed.text_content == "Recognized one\n\nRecognized two"
    assert parsed.parser_metadata["ocr_used"] is True
    assert adapter.calls == [1, 2]
    assert parsed.warnings == []


def test_pdf_parser_ocr_fills_only_empty_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(
        "ingestion.parsers.pdf.PdfReader",
        _pdf_reader_with_pages(["Real page one", ""]),
    )
    adapter = _StubOcrAdapter({2: "OCR page two"})

    parser = PdfParser(ocr_adapter=adapter)
    parsed = parser.parse(_source("doc-pdf", DocumentFormat.PDF), b"mixed pdf")

    assert parsed.text_content == "Real page one\n\nOCR page two"
    assert adapter.calls == [2]  # OCR invoked only for the empty page
    assert parsed.parser_metadata["ocr_used"] is True


def test_pdf_parser_skips_ocr_when_text_present(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(
        "ingestion.parsers.pdf.PdfReader",
        _pdf_reader_with_pages(["First page", "Second page"]),
    )
    adapter = _StubOcrAdapter({1: "should not be used"})

    parser = PdfParser(ocr_adapter=adapter)
    parsed = parser.parse(_source("doc-pdf", DocumentFormat.PDF), b"text pdf")

    assert adapter.calls == []
    assert parsed.parser_metadata["ocr_used"] is False


def test_pdf_parser_raises_for_image_pdf_without_adapter(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr("ingestion.parsers.pdf.PdfReader", _pdf_reader_with_pages(["", ""]))

    parser = PdfParser()  # no OCR adapter configured -> opt-in unchanged
    with pytest.raises(ParserError, match="does not contain extractable text"):
        parser.parse(_source("doc-pdf", DocumentFormat.PDF), b"image pdf")


def test_pdf_parser_keeps_empty_page_warning_when_ocr_finds_nothing(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        "ingestion.parsers.pdf.PdfReader",
        _pdf_reader_with_pages(["Real page one", ""]),
    )
    adapter = _StubOcrAdapter({})  # OCR recognizes nothing on the empty page

    parser = PdfParser(ocr_adapter=adapter)
    parsed = parser.parse(_source("doc-pdf", DocumentFormat.PDF), b"mixed pdf")

    assert adapter.calls == [2]
    assert parsed.parser_metadata["ocr_used"] is False
    assert any(w.code == "pdf.empty_page" and w.page_number == 2 for w in parsed.warnings)


@pytest.mark.skipif(
    importlib.util.find_spec("pdf2image") is not None,
    reason="OCR extra installed; covered by test_pdf_ocr_integration.py instead",
)
def test_tesseract_adapter_raises_clear_error_without_ocr_extra() -> None:
    from ingestion.parsers.adapters.tesseract import TesseractOcrAdapter

    adapter = TesseractOcrAdapter()
    with pytest.raises(ImportError, match="optional OCR dependencies"):
        adapter.recognize_page(b"pdf bytes", 1)


class TestHtmlUnclosedTagsDoNotTruncateTheDocument:
    """``html.parser`` performs no implicit tag closing.

    An unclosed ``<a>`` or ``<table>`` routes every later text node into a
    buffer that is never drained, so the tail of the document is dropped.
    Nothing surfaced the loss: the document was marked PARSED with an empty
    warnings list, and chunking, extraction and graph build all ran on the
    truncated text.
    """

    warnings: list[ParserWarning] = []

    def _parse(self, body: bytes) -> str:
        """Parse and return the rendered text, asserting it is present."""
        parsed = HtmlParser().parse(
            _source("doc-html-unclosed", DocumentFormat.HTML), body
        )
        self.warnings = parsed.warnings
        assert parsed.text_content is not None
        return parsed.text_content

    def test_text_after_an_unclosed_anchor_survives(self) -> None:
        text = self._parse(
            b"<html><body><p>Before the link.</p>"
            b'<a href="https://example.com">Link text'
            b"<p>NPI 1234567890 billed 42000 dollars.</p>"
            b"</body></html>"
        )

        assert "1234567890" in text
        assert "42000" in text

    def test_text_after_an_unclosed_table_survives(self) -> None:
        text = self._parse(
            b"<html><body><p>Before the table.</p>"
            b"<table><tr><td>Cell</td></tr>"
            b"<p>NPI 1234567890 billed 42000 dollars.</p>"
            b"</body></html>"
        )

        assert "1234567890" in text
        assert "42000" in text

    def test_an_unclosed_tag_is_reported_as_a_warning(self) -> None:
        """Truncation risk must reach DocumentsParsedEvent.warning_count."""
        self._parse(
            b"<html><body><p>Before.</p>"
            b"<table><tr><td>Cell</td></tr>"
            b"<p>After.</p></body></html>"
        )

        assert [w.code for w in self.warnings if w.code == "html.unclosed_tag"] == [
            "html.unclosed_tag"
        ]

    def test_well_formed_html_reports_no_unclosed_warning(self) -> None:
        text = self._parse(
            b"<html><body><p>Before.</p>"
            b"<table><tr><td>Cell</td></tr></table>"
            b"<p>After.</p></body></html>"
        )

        assert [w for w in self.warnings if w.code == "html.unclosed_tag"] == []
        assert "After." in text
